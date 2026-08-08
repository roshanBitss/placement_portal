
# Create your views here.
import traceback
from urllib import request
from django.shortcuts import render, redirect
from django.http import HttpResponse
from .models import Login, Quiz, UserSubmission
from django.contrib.auth.hashers import check_password, make_password
from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.template import loader
from .models import Profile
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import render, redirect
from django.http import HttpResponse
from .models import Login  # Ensure the Login model is correctly imported
from django.contrib.auth import logout
from django.http import HttpResponse
import os
from docx2pdf import convert
from docx import Document
import docx
from django.utils.timezone import now
from .models import Job,Event
from django.http import JsonResponse
import json
from twilio.rest import Client
from django.conf import settings
from credentials.models import Question
from credentials.serializers import QuestionSerializer
from django.shortcuts import render
from django.http import JsonResponse
from django.shortcuts import render
from django.http import JsonResponse
from django.contrib import messages
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from django.http import JsonResponse, HttpResponse
import os
import traceback
from django.http import FileResponse
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from django.http import JsonResponse, HttpResponse
import os
import traceback
from .models import Profile

def login_user(request):
    if request.method == "POST":
        username = request.POST.get("username")
        password = request.POST.get("password")

        # Check if the user exists in the database
        try:
            user = Login.objects.get(username=username)
            if password == user.password:  # Validate password
                #print("logged in")
                # Save username in session
                request.session['username'] = username
                #print(request.session['username']) 

                # Redirect to 'home' or another page
                return redirect("home/")
            else:
                messages.error(request, "Password is Invalid!")
        except Login.DoesNotExist:
            messages.error(request, "Username Does Not Exist!")
    # Render login page for GET requests
    return render(request, "myapp/login.html")

def profile_view(request):
    username = request.session.get('username')  # Get username from session
    if not username:
        return redirect("myapp/login.html")  # Redirect to login if not logged in

    # Fetch profile data (assuming a Profile model exists)
    try:
        profile = Profile.objects.get(contact__email=username)  # Adjust query as needed
        print(profile)
    except Profile.DoesNotExist:
        return HttpResponse("Profile not found!")

    return render(request, "profile.html", {
        "login_user": {"username": username},  # Pass username to the template
        "profile": profile
    })
    
#@login_required
def home(request):
        # Debug session data
    print("Session Data at Home:", request.session.items())
    
    # Check if the user is logged in
    session_username = request.session.get('username')
    if not session_username:
        return redirect('')  # Redirect to login if session is empty
    
    # Fetch the user's profile
    try:
        profile = Profile.objects.get(username__username=session_username)
    except Profile.DoesNotExist:
        profile = None  
        
    data_points = [
        {"label": "2024", "y": 8},
        {"label": "2023", "y": 2},
        {"label": "2022", "y": 3},
        {"label": "2021", "y": 4},
        {"label": "2020", "y": 5},
        {"label": "2019", "y": 6},
    ]
      # Data points for Sales Channels
    data_points_sales = [
        {"label": "BAF", "y": 27},
        {"label": "BMS", "y": 25},
        {"label": "Bsc IT", "y": 30},
        {"label": "BMM", "y": 8},
        {"label": "Others", "y": 10},
    ]
    
    
   # Pass both datasets to the template
    return render(request, "myapp/home.html", {
        "data_points_fruits": data_points,
        "data_points_sales": data_points_sales,
        "profile": profile,
    })
    

def profile(request):
    # Get the username from the session
    session_username = request.session.get('username')
    
    if session_username:
        try:
            # Fetch the profile of the logged-in user using `username`
            profile_data = Profile.objects.prefetch_related('experiences').select_related('contact').get(username__username=session_username)

            context = {
                "profile": profile_data,
                "projects": profile_data.projects.split(",") if profile_data.projects else [],
                "skills": profile_data.skills.split(",") if profile_data.skills else [],
                "education": profile_data.education.split(",") if profile_data.education else [],
                "experience": profile_data.experiences.all(),
                "interests": profile_data.interests.split(",") if profile_data.interests else [],
                "achievements": profile_data.achievements.split(",") if profile_data.achievements else [],
            }
        except Profile.DoesNotExist:
            context = {"error_message": "Profile not found for the user."}
    else:
        # If session is empty, redirect to login
        return redirect("")

    return render(request, 'myapp/profile.html', context)

def update_about_me(request):
    if request.method == "POST":
        username = request.POST.get("username")
        about_me = request.POST.get("about_me")

        if not username or not about_me:
            return JsonResponse({"status": "error", "error": "Invalid data"})

        try:
            profile = Profile.objects.get(contact__email=username)  # Adjust query if needed
            profile.about_me = about_me
            profile.save()
            return JsonResponse({"status": "success"})
        except Profile.DoesNotExist:
            return JsonResponse({"status": "error", "error": "Profile not found"})
    return JsonResponse({"status": "error", "error": "Invalid request method"})

def job(request):
    session_username = request.session.get('username')
    if not session_username:
        return redirect('')  # Redirect to login if not logged in

    try:
        profile = Profile.objects.get(username__username=session_username)
    except Profile.DoesNotExist:
        profile = None
    
    jobs = Job.objects.filter(last_date_for_submission__gte=now()).order_by('last_date_for_submission')
     # Fetch available jobs
    jobs = Job.objects.filter(last_date_for_submission__gte=now()).order_by('last_date_for_submission')

    # Match jobs with the user's skills or role
    matched_jobs = []
    for job in jobs:
        job_skills = job.skills_required.split(",") if job.skills_required else []
        profile_skills = profile.skills.split(",") if profile.skills else []
        
        # Check for any match in skills or role
        if set(job_skills).intersection(profile_skills) or profile.role == job.job_role:
            matched_jobs.append(job)

    # Notify user via WhatsApp for matched jobs
    if matched_jobs:
        client = Client(settings.TWILIO_ACCOUNT_SID, settings.TWILIO_AUTH_TOKEN)
        for job in matched_jobs:
            message_body = (
                f"Hi {profile.name},\n\n"
                f"We found a job that matches your profile!\n\n"
                f"Job Title: {job.job_role}\n"
                f"Company: {job.company_name}\n"
                f"Skills Required: {job.skills_required}\n"
                f"Last Date to Apply: {job.last_date_for_submission.strftime('%d-%m-%Y')}\n\n"
                f"Apply Now: {job.application_link}\n\n"
                f"Apply soon!"
            )

            try:
                # Send WhatsApp message
                client.messages.create(
                    body=message_body,
                    from_=settings.TWILIO_WHATSAPP_FROM,
                    to=f'whatsapp:{profile.contact.phone.replace(" ", "")}'  # User's phone number
                )
                print(f'Message sent successfully!{profile.contact.phone.replace(" ", "")}')
            except Exception as e:
                print(f"Failed to send WhatsApp message: {e}")

    return render(request, "myapp/job.html", {
        "profile": profile,
        "jobs": jobs,
        "matched_jobs": matched_jobs
    })

    

def event(request):
    session_username = request.session.get('username')
    if not session_username:
        return redirect('')  # Redirect to login if not logged in

    try:
        profile = Profile.objects.get(username__username=session_username)
    except Profile.DoesNotExist:
        profile = None
        
    events = Event.objects.filter(event_date__gte=now()).order_by('event_date')

    return render(request, "myapp/event.html", {
        "profile": profile,
        "events":events
    })
    
    
def logout_user(request):
    # Log out the user
    logout(request)
    
    # Redirect to the login page after logout
    return redirect('login_user')  # Replace 'login' with the name of your login URL

def generate_resume(request):
    session_username = request.session.get('username')
    if not session_username:
        return redirect('login')

    try:
        profile = Profile.objects.get(username__username=session_username)
    except Profile.DoesNotExist:
        return HttpResponse("Profile not found.", status=404)

    doc = Document()

    # Personal Information
    doc.add_heading("Personal Information", level=1)
    doc.add_paragraph(f"Name: {profile.name} | Email: {profile.contact.email} | Phone: {profile.contact.phone}\nLocation: {profile.contact.location}")
    doc.add_paragraph("\n" + "_" * 100)
    # About Me
    doc.add_heading("About Me", level=2)
    doc.add_paragraph(profile.about_me)
    doc.add_paragraph("\n" + "_" * 100)
    # Education
    doc.add_heading("Education", level=2)
    if profile.education:
        for edu in profile.education.split(","):
            doc.add_paragraph(f"- {edu.strip()}", style="List Bullet")
    doc.add_paragraph("\n" + "_" * 100)
    # Skills & Experience
    doc.add_heading("Skills & Experience", level=2)
    if profile.skills:
        doc.add_paragraph("Skills:")
        for skill in profile.skills.split(","):
            doc.add_paragraph(f"- {skill.strip()}", style="List Bullet")
    if profile.experiences.exists():
        doc.add_paragraph("Experience:")
        for exp in profile.experiences.all():
            doc.add_paragraph(f"- {exp.title}: {exp.description}", style="List Bullet")
    doc.add_paragraph("\n" + "_" * 100)
    # Projects
    doc.add_heading("Projects", level=2)
    if profile.projects:
        for project in profile.projects.split(","):
            doc.add_paragraph(f"- {project.strip()}", style="List Bullet")
    doc.add_paragraph("\n" + "_" * 100)
    # Achievements
    doc.add_heading("Achievements", level=2)
    if profile.achievements:
        for achievement in profile.achievements.split(","):
            doc.add_paragraph(f"- {achievement.strip()}", style="List Bullet")
    
    # Save the Word document to a temporary file
    docx_file_path = "resume.docx"
    doc.save(docx_file_path)

    # Convert to PDF
    try:
        pdf_file_path = "resume.pdf"
        convert(docx_file_path, pdf_file_path)
    except Exception as e:
        return HttpResponse(f"Error converting to PDF: {str(e)}", status=500)

    with open(pdf_file_path, 'rb') as pdf_file:
        pdf_content = pdf_file.read()

    os.remove(docx_file_path)
    os.remove(pdf_file_path)

    response = HttpResponse(pdf_content, content_type="application/pdf")
    response["Content-Disposition"] =f'attachment; filename="{profile.name}resume.pdf"'
    return response


def update_profile_image(request):
    if request.method == 'POST':
        profile = Profile.objects.get(user=request.user)  # Fetch the logged-in user's profile
        if 'profile_image' in request.FILES:
            profile.image = request.FILES['profile_image']  # Updates the image field on the basis of users credentials 
            profile.save()  
        return redirect('profile')  # Redirect to the home page or profile page
    
@csrf_exempt  # Required for AJAX post requests (you can also use authentication if needed)
def save_contact_details(request):
    if request.method == 'POST':
        # Check if the user is authenticated via session
        if request.session.get('user_id'):
            user_id = request.session.get('user_id')
            user = Login.objects.get(id=user_id)  # Get the User object using session data

            # Parse the JSON data from the request body
            data = json.loads(request.body)
            email = data.get('email')
            phone = data.get('phone')
            location = data.get('location')

            try:
                # Fetch the profile of the logged-in user
                profile = Profile.objects.get(user=user)

                # Update the contact details
                profile.contact.email = email
                profile.contact.phone = phone
                profile.contact.location = location
                profile.save()

                return JsonResponse({'success': True})

            except Profile.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Profile not found'}, status=404)
        else:
            return JsonResponse({'success': False, 'error': 'User is not authenticated'}, status=401)

    return JsonResponse({'success': False, 'error': 'Invalid request method'}, status=400)

def quiz(request):
    session_username = request.session.get('username')

    if not session_username:
        return redirect('login')

    try:
        profile = Profile.objects.filter(username__username=session_username).first()
    except Profile.DoesNotExist:
        profile = None

    return render(request, 'myapp/test.html', {'profile': profile})


@csrf_exempt

def questionsViews(request, username, quiz_type):
    session_username = request.session.get('username')
    if not session_username:
        return redirect('login')

    try:
        profile = Profile.objects.filter(username__username=session_username).first()
        user = Login.objects.get(username=session_username)
    except (Profile.DoesNotExist, Login.DoesNotExist):
        profile = None
        user = None

    # Initialize session storage for quiz answers
    if 'quiz_answers' not in request.session:
        request.session['quiz_answers'] = {}

    if request.method == "POST":
        try:
            data = json.loads(request.body)
            question_id = str(data.get("question_id"))  # Ensure question_id is a string
            selected_answer = data.get("selected_answer", "").strip().lower()

            if not question_id or not selected_answer:
                return JsonResponse({"error": "Missing question_id or selected_answer"}, status=400)

            question = Question.objects.filter(id=question_id, category__iexact=quiz_type).first()
            if not question:
                return JsonResponse({"error": "Question not found"}, status=400)

           # answer keys
            answer_map = {
                "a": question.option_a.strip().lower(),
                "b": question.option_b.strip().lower(),
                "c": question.option_c.strip().lower(),
                "d": question.option_d.strip().lower(),
            }

            correct_answer_key = question.correct_answer.strip().lower()
            correct_answer_text = answer_map.get(correct_answer_key, "")

            #  Comparing the answer correctly
            is_correct = selected_answer == correct_answer_key

            # Store answer temporarily in session
            quiz_answers = request.session.get('quiz_answers', {})
            if quiz_type not in quiz_answers:
                quiz_answers[quiz_type] = {}

            quiz_answers[quiz_type][question_id] = selected_answer
            request.session['quiz_answers'] = quiz_answers
            request.session.modified = True

            # Checking if all questions are answered
            total_questions = Question.objects.filter(category__iexact=quiz_type).count()
            answered_questions = len(quiz_answers[quiz_type])

            print(f"DEBUG: Answered {answered_questions}/{total_questions}")

            if answered_questions == total_questions:
                print("DEBUG: All questions answered, calculating score...")

                #  Calculating score
                score = 0
                for q_id, ans in quiz_answers[quiz_type].items():
                    question = Question.objects.filter(id=int(q_id)).first()
                    if question:
                        print(f"DEBUG: Checking Question ID {q_id}: Correct Answer = {question.correct_answer.strip().lower()}, User Answer = {ans}")
                        if question.correct_answer.strip().lower() == ans:
                            score += 1
                    else:
                        print(f"DEBUG ERROR: Question ID {q_id} not found!")

                # Ensuring user exists before submission
                if not user:
                    print("DEBUG ERROR: User is None!")
                    return JsonResponse({"error": "User not found"}, status=400)

                #  Saving final submission
                category_name = question.category
                
                UserSubmission.objects.create(
                    username=user,
                    category=category_name,
                    score=score,
                    submitted_at=now()
                )

                #  Clearing the session data for this quiz
                if quiz_type in request.session['quiz_answers']:
                    del request.session['quiz_answers'][quiz_type]
                    request.session.modified = True

                return JsonResponse({
                    "message": "Quiz completed and submitted!",
                    "total_score": score
                })

            return JsonResponse({
                "message": "Answer saved!",
                "is_correct": is_correct,
                "correct_answer": correct_answer_text,
                "answered_questions": answered_questions,
                "total_questions": total_questions
            })

        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON format"}, status=400)
        except Exception as e:
            print(f"DEBUG ERROR: {str(e)}")
            return JsonResponse({"error": str(e)}, status=500)

    elif request.method == "GET":
        try:
            questions = Question.objects.filter(category__iexact=quiz_type).values(
                "id", "question_text", "option_a", "option_b", "option_c", "option_d"
            )
            questions_list = list(questions)

            for question in questions_list:
                question["options"] = [
                    question.pop("option_a"),
                    question.pop("option_b"),
                    question.pop("option_c"),
                    question.pop("option_d"),
                ]
            
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':  
                return JsonResponse({"questions": questions_list})
            
            return render(request, 'myapp/questions.html', {
                'username': username,
                'quiz_type': quiz_type,
                'profile': profile,
                'questions': json.dumps(questions_list)
            })
        except Exception as e:
            print(f"DEBUG ERROR: {str(e)}")
            return JsonResponse({"status": "error", "message": str(e)}, status=500)

     #Resume making
def resume_view(request):
    profile = Profile.objects.get(username=request.user)  # Fetch logged-in user's profile
    skills_list = profile.skills.split(",") if profile.skills else []
    projects_list = profile.projects.split(",") if profile.projects else []
    education_list = profile.education.split(",") if profile.education else []
    interests_list = profile.interests.split(",") if profile.interests else []
    achievements_list = profile.achievements.split(",") if profile.achievements else []

    context = {
        'profile': profile,
        'skills': skills_list,
        'projects': projects_list,
        'education': education_list,
        'interests': interests_list,
        'achievements': achievements_list
    }
    return render(request, 'resume.html', context)